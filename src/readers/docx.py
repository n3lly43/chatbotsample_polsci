"""Reader for Word .docx files."""


def read_docx(file_path: str) -> list[dict]:
    from docx import Document
    doc = Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Extract embedded tables
    for table in doc.tables:
        for row in table.rows:
            seen_elements = set()
            cells = []
            for cell in row.cells:
                if id(cell._element) not in seen_elements:
                    seen_elements.add(id(cell._element))
                    if cell.text.strip():
                        cells.append(cell.text.strip())
            if cells:
                paragraphs.append(" | ".join(cells))

    if not paragraphs:
        return []
    return [{"page": 1, "text": "\n\n".join(paragraphs)}]
