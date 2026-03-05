import pytest

def test_docx_reader(tmp_path):
    pytest.importorskip("docx")
    from docx import Document
    doc = Document()
    doc.add_paragraph("First paragraph of the document.")
    doc.add_paragraph("Second paragraph with more content.")
    path = tmp_path / "test.docx"
    doc.save(str(path))
    from src.readers.docx import read_docx
    pages = read_docx(str(path))
    assert len(pages) >= 1
    assert "First paragraph" in pages[0]["text"]

def test_xlsx_reader(tmp_path):
    pytest.importorskip("openpyxl")
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(["Country", "Year", "Success"])
    ws.append(["India", 1930, "Yes"])
    ws.append(["Poland", 1980, "Yes"])
    path = tmp_path / "test.xlsx"
    wb.save(str(path))
    from src.readers.excel import read_excel
    pages = read_excel(str(path))
    assert len(pages) >= 1
    assert "India" in pages[0]["text"]
    assert "Country" in pages[0]["text"]

def test_csv_reader(tmp_path):
    f = tmp_path / "data.csv"
    f.write_text("name,year,outcome\nAlice,2020,success\nBob,2021,failure\n")
    from src.readers.csv_tab import read_csv_tab
    pages = read_csv_tab(str(f))
    assert len(pages) >= 1
    assert "Alice" in pages[0]["text"]

def test_tab_reader(tmp_path):
    f = tmp_path / "data.tab"
    f.write_text("country\tyear\tresult\nEgypt\t2011\tpartial\n")
    from src.readers.csv_tab import read_csv_tab
    pages = read_csv_tab(str(f), delimiter="\t")
    assert len(pages) >= 1
    assert "Egypt" in pages[0]["text"]
